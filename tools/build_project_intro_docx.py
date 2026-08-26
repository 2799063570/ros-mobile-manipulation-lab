from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


OUT_DIR = Path(__file__).resolve().parents[1] / "deliverables"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "AUBO复合移动机器人项目介绍.docx"

NAVY = "17365D"
BLUE = "2F75B5"
TEAL = "1696A7"
ORANGE = "ED7D31"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F5F6"
LIGHT_ORANGE = "FCEFE5"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "667085"
DARK = "1F2937"
WHITE = "FFFFFF"
BORDER = "D7DEE7"
YELLOW = "FFF2CC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "start" if edge == "start" else "end" if edge == "end" else edge
            elem = borders.find(qn(f"w:{tag}"))
            if elem is None:
                elem = OxmlElement(f"w:{tag}")
                borders.append(elem)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    elem.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_geometry(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = int(sum(widths_in) * 1440)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(width * 1440)))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width_dxa = int(widths_in[min(idx, len(widths_in)-1)] * 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=None, bold=None, color=None, name="Microsoft YaHei", italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, 9, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run2 = paragraph.add_run(" 页")
    set_font(run2, 9, color=MID_GRAY)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Microsoft YaHei")
    r_fonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.extend([r_fonts, c, u])
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_para(doc, text="", size=10.5, bold=False, color=DARK, align=None,
             after=6, before=0, line=1.25, italic=False, keep=False):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        set_font(r, size, bold, color, italic=italic)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if keep:
        keep_with_next(p)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_font(r, 10.2, color=DARK)
    return p


def add_heading(doc, text, level=1, number=None):
    p = doc.add_paragraph(style=f"Heading {level}")
    label = f"{number}  {text}" if number else text
    r = p.add_run(label)
    set_font(r, 16 if level == 1 else 12.5, True, NAVY if level == 1 else BLUE)
    p.paragraph_format.keep_with_next = True
    return p


def add_section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text.upper())
    set_font(r, 8.5, True, TEAL)
    keep_with_next(p)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.27])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    edge = {"val": "single", "sz": "16", "space": "0", "color": accent}
    none = {"val": "nil"}
    set_cell_border(cell, start=edge, top=none, bottom=none, end=none)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, 10.5, True, accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_font(r2, 9.8, color=DARK)
    add_para(doc, "", after=2)
    return table


def add_feature_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        set_font(r, 9.5, True, WHITE)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], LIGHT_GRAY)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(value)
            set_font(r, 9.1, bold=(i == 0), color=DARK)
    border = {"val": "single", "sz": "4", "space": "0", "color": BORDER}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)
    return table


def add_process(doc, steps, fills=None):
    fills = fills or [LIGHT_BLUE] * len(steps)
    widths = [6.27 / len(steps)] * len(steps)
    table = doc.add_table(rows=1, cols=len(steps))
    set_table_geometry(table, widths)
    for i, (title, desc) in enumerate(steps):
        cell = table.cell(0, i)
        set_cell_shading(cell, fills[i % len(fills)])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{i+1:02d}")
        set_font(r, 9, True, TEAL)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(title)
        set_font(r2, 9.4, True, NAVY)
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(desc)
        set_font(r3, 7.8, color=MID_GRAY)
        b = {"val": "single", "sz": "4", "space": "0", "color": WHITE}
        set_cell_border(cell, top=b, bottom=b, start=b, end=b)
    return table


def add_page_title(doc, index, title, subtitle):
    add_section_label(doc, f"PROJECT BRIEF / {index:02d}")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, 22, True, NAVY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    set_font(r2, 10.5, color=MID_GRAY)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)       # Named override: A4 Chinese submission format
section.page_height = Cm(29.7)
section.top_margin = Cm(1.9)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.header_distance = Cm(0.9)
section.footer_distance = Cm(0.9)

# Global styles: standard_business_brief, adapted to Chinese typography and A4.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Heading 1", 16, NAVY, 14, 7),
    ("Heading 2", 12.5, BLUE, 10, 5),
    ("Heading 3", 11, TEAL, 8, 4),
]:
    st = styles[name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ["List Bullet", "List Bullet 2"]:
    st = styles[list_name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(10.2)

# Running header/footer.
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("AUBO 复合移动机器人项目介绍")
set_font(hr, 8.5, True, MID_GRAY)
footer = section.footer
fp = footer.paragraphs[0]
add_page_number(fp)

# PAGE 1 — editorial cover.
add_para(doc, "机器人系统集成 · 项目作品", size=9.5, bold=True, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, after=70)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run("基于 ROS 的 AUBO\n复合移动机器人")
set_font(r, 29, True, NAVY)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(22)
r2 = p2.add_run("自主导航、视觉识别、机械臂抓取与颜色分拣系统")
set_font(r2, 14, True, BLUE)
add_process(doc, [
    ("建图定位", "双雷达 / GMapping / AMCL"),
    ("自主导航", "move_base / DWA"),
    ("视觉感知", "手部相机 / OpenCV"),
    ("抓取分拣", "MoveIt / 轨迹控制"),
], [LIGHT_BLUE, LIGHT_TEAL, LIGHT_ORANGE, LIGHT_BLUE])
add_para(doc, "", after=18)
meta = doc.add_table(rows=4, cols=2)
set_table_geometry(meta, [1.35, 4.92])
metadata = [
    ("项目类型", "机器人仿真与系统集成项目"),
    ("开发环境", "Ubuntu 18.04 · ROS 1 Melodic · Gazebo Classic"),
    ("项目仓库", "github.com/2799063570/wheeltec_robot"),
    ("作者/团队", "请在投递前填写姓名、学校及团队成员"),
]
for i, (label, value) in enumerate(metadata):
    c0, c1 = meta.rows[i].cells
    set_cell_shading(c0, NAVY)
    set_cell_shading(c1, LIGHT_GRAY if i != 3 else YELLOW)
    for c in (c0, c1):
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        border = {"val": "single", "sz": "4", "space": "0", "color": WHITE}
        set_cell_border(c, top=border, bottom=border, start=border, end=border)
    p0 = c0.paragraphs[0]; p0.paragraph_format.space_after = Pt(0)
    rr0 = p0.add_run(label); set_font(rr0, 9.2, True, WHITE)
    p1 = c1.paragraphs[0]; p1.paragraph_format.space_after = Pt(0)
    rr1 = p1.add_run(value); set_font(rr1, 9.2, bold=(i == 2), color=DARK)
add_para(doc, "项目作品介绍 / 求职投递版", size=9, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, before=28, after=0)

# PAGE 2 — summary.
doc.add_page_break()
add_page_title(doc, 2, "项目概览", "从移动机器人学习验证，逐步完成复合移动操作系统集成")
add_callout(doc, "一句话介绍", "面向室内作业场景，系统能够先完成环境建图与自主导航，再利用手部相机识别红、绿、蓝目标，通过 AUBO i5 六轴机械臂完成抓取和分类放置。", LIGHT_BLUE, BLUE)
add_heading(doc, "项目目标", 2)
add_para(doc, "本项目围绕“移动到达 + 环境感知 + 机械操作”这一复合机器人任务链展开。我们先以 WheelTec ROS 功能包为学习参考，理解底盘、传感器、TF、SLAM 与导航的数据链路；随后自主搭建差速移动机器人，并将 AUBO i5、双指夹爪、双激光雷达和手部 RGB 相机集成到统一的 ROS 系统中。")
add_heading(doc, "核心能力", 2)
add_feature_table(doc, ["能力模块", "主要实现", "输出效果"], [
    ("机器人建模与仿真", "Xacro/URDF、Gazebo 传感器与 ros_control 控制器", "形成可移动、可感知、可规划的复合机器人模型"),
    ("建图与定位", "前后激光雷达融合、GMapping、AMCL、TF 管理", "生成地图并获得 map → odom → base_link 位姿链"),
    ("自主导航", "move_base、全局/局部代价地图、DWA 局部规划", "根据工位目标自主规划并驱动底盘到达"),
    ("视觉感知", "HSV 颜色分割、轮廓检测、像素射线与平面求交", "输出目标颜色、像素坐标与 base_link 下三维位置"),
    ("机械臂操作", "MoveIt、笛卡尔路径、轨迹控制器与夹爪控制", "完成预抓取、下降、夹取、抬升和分类放置"),
    ("任务编排", "状态机、ROS 服务/话题、导航与分拣顺序协调", "实现一键启动的“导航到工位后自动分拣”流程"),
], [1.35, 2.72, 2.20])
p_kw = doc.add_paragraph()
p_kw.paragraph_format.space_before = Pt(7)
p_kw.paragraph_format.space_after = Pt(0)
p_kw.paragraph_format.line_spacing = 1.05
r_kw1 = p_kw.add_run("技术关键词：")
set_font(r_kw1, 9.4, True, BLUE)
r_kw2 = p_kw.add_run("ROS 1 / C++ / Python / Gazebo / RViz / MoveIt / OpenCV / GMapping / AMCL / Navigation Stack / TF")
set_font(r_kw2, 9.2, True, TEAL)

# PAGE 3 — evolution and contribution boundary.
doc.add_page_break()
add_page_title(doc, 3, "项目演进与成果边界", "清晰呈现学习来源、自主实现和系统集成贡献")
add_process(doc, [
    ("开源学习", "WheelTec 功能包\n理解 ROS 数据链"),
    ("差速机器人", "自主模型\n建图定位导航"),
    ("AUBO 机械臂", "MoveIt 规划\n抓取放置验证"),
    ("复合机器人", "导航 + 视觉 + 分拣\n完整任务闭环"),
], [LIGHT_GRAY, LIGHT_BLUE, LIGHT_TEAL, LIGHT_ORANGE])
add_heading(doc, "主要实践成果", 2)
add_feature_table(doc, ["成果目录", "定位", "代表性内容"], [
    ("simple_diff_robot_gazebo", "自主差速移动机器人", "模型、差速驱动、激光与相机、建图、AMCL 导航、RRT 探索"),
    ("aubo/aubo_planning", "机械臂规划实践", "夹爪控制、MoveIt 抓取与放置流程"),
    ("aubo_mobile_robot/", "复合机器人主成果", "模型、导航、感知、控制、分拣和场景任务编排"),
], [1.80, 1.72, 2.75])
add_heading(doc, "开源参考与自主工作的区分", 2)
add_callout(doc, "投递说明", "仓库保留 WheelTec、Navigation、MoveIt 等第三方或学习参考功能包，用于理解成熟 ROS 系统的组织方式。项目介绍不将这些内容表述为完全自研；重点展示在其基础上完成的模型搭建、功能整合、参数配置、任务编排和系统联调。", LIGHT_ORANGE, ORANGE)
add_heading(doc, "建议填写的个人贡献", 2)
add_para(doc, "请根据实际情况保留并修改以下内容，避免夸大：", size=9.8, color=MID_GRAY)
for item in [
    "负责复合机器人 Xacro/URDF 模型、传感器与控制器集成；",
    "负责建图、AMCL 定位、move_base 导航参数配置与 TF 排错；",
    "负责 OpenCV 颜色目标识别、平面目标定位与标定参数设计；",
    "负责 MoveIt 抓放流程、夹爪控制、PlanningScene 碰撞物体和任务状态机；",
    "负责导航、视觉和分拣模块联调、异常保护与文档整理。",
]:
    add_bullet(doc, item)

# PAGE 4 — architecture.
doc.add_page_break()
add_page_title(doc, 4, "系统组成与软件架构", "底盘导航与机械臂操作分层实现，通过任务协调器形成闭环")
add_heading(doc, "硬件/仿真组成", 2)
add_feature_table(doc, ["组成", "配置", "作用"], [
    ("移动底盘", "差速驱动底盘 + 编码器里程计", "执行速度指令并发布 /odom"),
    ("机械臂", "AUBO i5 六轴机械臂", "执行 MoveIt 规划后的关节轨迹"),
    ("末端执行器", "双指夹爪 + tcp_link", "完成方块夹取和释放"),
    ("环境感知", "前后二维激光雷达", "融合生成 360° /scan，用于建图、定位和避障"),
    ("操作感知", "手部 RGB 相机", "识别颜色目标并估计桌面位置"),
], [1.35, 2.20, 2.72])
add_heading(doc, "ROS 数据流", 2)
arch = doc.add_table(rows=5, cols=3)
set_table_geometry(arch, [1.65, 2.85, 1.77])
arch_rows = [
    ("传感器层", "/front/scan + /rear/scan", "双雷达融合"),
    ("导航层", "/scan + /odom + TF", "GMapping / AMCL / move_base"),
    ("感知层", "/hand_camera/image_raw", "OpenCV 颜色识别与目标定位"),
    ("规划控制层", "目标位姿 + PlanningScene", "MoveIt + 机械臂/夹爪控制器"),
    ("任务层", "/nav_sorting/state", "导航、观察、分拣状态机"),
]
for ridx, row in enumerate(arch_rows):
    for cidx, text in enumerate(row):
        cell = arch.cell(ridx, cidx)
        set_cell_shading(cell, [NAVY, LIGHT_BLUE, LIGHT_TEAL][cidx] if cidx > 0 else NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        rr = p.add_run(text)
        set_font(rr, 9.2, bold=(cidx != 1), color=WHITE if cidx == 0 else DARK)
        border = {"val": "single", "sz": "5", "space": "0", "color": WHITE}
        set_cell_border(cell, top=border, bottom=border, start=border, end=border)
add_heading(doc, "关键坐标链", 2)
add_callout(doc, "TF 主链路", "map → odom → base_footprint → base_link → AUBO 各级连杆 → tcp_link", LIGHT_TEAL, TEAL)
add_para(doc, "建图时 map → odom 由 GMapping 发布；使用已有地图导航时由 AMCL 发布。二者不能同时运行，否则会产生 TF 冲突。", size=9.8, color=MID_GRAY)

# PAGE 5 — navigation.
doc.add_page_break()
add_page_title(doc, 5, "移动导航子系统", "从双雷达扫描到工位目标到达，构成可复用的移动能力")
add_heading(doc, "双雷达融合", 2)
add_para(doc, "前、后激光扫描首先变换到 base_footprint 坐标系，再生成包含 720 个采样点的 360° /scan。融合过程考虑雷达安装位置和朝向，并非简单拼接距离数组。")
add_process(doc, [
    ("采集", "前后 LaserScan"),
    ("坐标变换", "转换到 base_footprint"),
    ("扫描融合", "生成 360° /scan"),
    ("导航消费", "SLAM / AMCL / move_base"),
], [LIGHT_BLUE, LIGHT_TEAL, LIGHT_BLUE, LIGHT_ORANGE])
add_heading(doc, "建图、定位与导航", 2)
add_feature_table(doc, ["阶段", "核心模块", "结果"], [
    ("环境建图", "GMapping + 里程计 + 融合激光", "生成并保存 map.yaml / map.pgm"),
    ("初始定位", "AMCL + 静态地图", "估计机器人在地图中的概率位姿"),
    ("路径规划", "move_base + 全局规划器 + DWA", "生成全局路径并进行局部避障"),
    ("底盘执行", "/cmd_vel + 差速驱动", "跟踪速度指令并更新 /odom"),
], [1.30, 2.50, 2.47])
add_heading(doc, "工程安全约束", 2)
for item in [
    "导航 footprint 按机械臂收拢状态设置，底盘移动前先将机械臂收回 home 或 down 安全姿态；",
    "机械臂伸展或抓取期间保持底盘静止，避免感知坐标和碰撞边界失效；",
    "仿真与实机切换时分别检查里程计、TF 发布者和控制器，避免重复发布坐标变换。",
]:
    add_bullet(doc, item)
add_callout(doc, "导航输出", "机器人在完成定位后接收工位目标位姿，经 move_base 规划和底盘执行，到达机械臂可覆盖工作台的预设作业位置。", LIGHT_BLUE, BLUE)

# PAGE 6 — perception and manipulation.
doc.add_page_break()
add_page_title(doc, 6, "视觉感知与抓取分拣", "利用手部 RGB 相机定位桌面目标，通过 MoveIt 完成抓放动作")
add_heading(doc, "颜色目标定位", 2)
add_process(doc, [
    ("图像输入", "手部 RGB 相机"),
    ("颜色分割", "HSV 阈值"),
    ("目标筛选", "轮廓面积 / 长宽比"),
    ("位置估计", "像素射线与桌面求交"),
    ("坐标发布", "base_link 下三维位置"),
], [LIGHT_BLUE, LIGHT_TEAL, LIGHT_BLUE, LIGHT_TEAL, LIGHT_ORANGE])
add_para(doc, "当前方案面向高度已知、表面平整的桌面场景。节点根据相机内参生成像素射线，并在 base_link 坐标系下与标定平面求交；连续 8 帧同色目标坐标求平均，以降低单帧轮廓中心抖动。")
add_heading(doc, "单个目标的抓放流程", 2)
add_process(doc, [
    ("观察", "获取最新识别坐标"),
    ("预抓取", "张开夹爪并规划到目标上方"),
    ("抓取", "笛卡尔下降并闭合夹爪"),
    ("转运", "抬升并移动到分类区域"),
    ("放置", "下降、释放并安全回退"),
], [LIGHT_TEAL, LIGHT_BLUE, LIGHT_ORANGE, LIGHT_BLUE, LIGHT_TEAL])
add_heading(doc, "关键工程处理", 2)
add_feature_table(doc, ["问题", "处理方式"], [
    ("桌面碰撞", "在机械臂运动前将 sorting_table 加入 MoveIt PlanningScene，并等待场景确认"),
    ("识别抖动", "对连续多帧目标位置求平均，并保留标定偏移参数"),
    ("夹爪接触失败", "调整夹爪闭合目标、运动时间及接触容差，符合实体接触特性"),
    ("小方块易滑落", "Gazebo 中使用临时固定关节提高仿真稳定性；实机不依赖该插件"),
    ("关节越界风险", "统一 URDF 与 MoveIt 中 upperArm_joint 的关节限制并在启动时校验"),
], [1.62, 4.65])

# PAGE 7 — mission and validation.
doc.add_page_break()
add_page_title(doc, 7, "完整任务编排与验证", "以状态机连接移动、感知和操作模块，形成端到端任务闭环")
add_heading(doc, "导航分拣任务流程", 2)
add_process(doc, [
    ("收臂", "STOWING_ARM"),
    ("导航", "NAVIGATING"),
    ("到达", "AT_WORKSTATION"),
    ("观察分拣", "SORTING"),
    ("完成", "SUCCEEDED"),
], [LIGHT_GRAY, LIGHT_BLUE, LIGHT_TEAL, LIGHT_ORANGE, LIGHT_TEAL])
add_para(doc, "任务协调器不重复实现底层算法，而是复用导航、MoveIt、视觉识别和分拣模块，通过服务和状态话题完成启动、停止、超时、失败与结果发布。")
add_heading(doc, "主要接口", 2)
add_feature_table(doc, ["接口", "作用"], [
    ("/cmd_vel", "底盘速度指令"),
    ("/odom / /scan", "轮式里程计与融合激光扫描"),
    ("/move_base", "底盘导航 action"),
    ("/sorting/detections", "颜色目标及三维位置"),
    ("/sorting/state", "视觉分拣子任务状态"),
    ("/nav_sorting/state", "导航分拣总任务状态"),
    ("follow_joint_trajectory", "机械臂和夹爪轨迹执行接口"),
], [2.25, 4.02])
add_heading(doc, "验证方式", 2)
for item in [
    "在 Gazebo 中验证机器人模型、传感器、控制器及物理接触；",
    "在 RViz 中检查 TF、地图、激光、代价地图、路径和 MoveIt 规划场景；",
    "通过调试图像核对目标轮廓、颜色类别、像素中心和定位结果；",
    "通过状态话题和 ROS 服务验证任务启动、停止、异常进入 ERROR/FAILED 和恢复逻辑。",
]:
    add_bullet(doc, item)
add_callout(doc, "当前实现范围", "采用“先移动、后操作”的顺序协调方案，尚未实现底盘与机械臂同时参与的全身运动规划。该设计更适合教学验证和室内固定工位任务，也便于分模块调试。", LIGHT_ORANGE, ORANGE)

# PAGE 8 — conclusion and link.
doc.add_page_break()
add_page_title(doc, 8, "项目价值、局限与后续方向", "以完整系统能力展示机器人软件开发、集成和调试经验")
add_heading(doc, "项目亮点", 2)
for item in [
    "完整性：覆盖模型、传感器、SLAM、定位、导航、视觉、规划、控制和任务编排；",
    "模块化：导航、感知、分拣与场景任务分离，可独立启动和调试；",
    "工程意识：处理 TF 冲突、碰撞场景、关节限制、夹爪接触容差和任务失败状态；",
    "可迁移性：仿真入口与实机入口分离，关键标定项和安全约束均有明确说明。",
]:
    add_bullet(doc, item)
add_heading(doc, "已知局限", 2)
add_para(doc, "仿真参数不能直接替代实机标定；RGB 单目定位依赖已知桌面高度；导航 footprint 仅覆盖机械臂收拢状态；Gazebo 抓取稳定插件仅服务于仿真验证。实机部署还需要重新完成地图、雷达外参、相机内外参、工位位姿、TCP、夹爪行程、速度限制、急停与碰撞保护配置。")
add_heading(doc, "可继续拓展", 2)
add_feature_table(doc, ["方向", "可实施内容"], [
    ("三维感知", "引入深度相机或点云，替代固定平面求交定位"),
    ("抓取鲁棒性", "加入姿态估计、抓取评分、力/触觉反馈和失败重试"),
    ("导航能力", "接入动态避障、语义地图与多工位任务调度"),
    ("复合规划", "研究底盘与机械臂联合规划及移动操作全身控制"),
    ("实机工程化", "统一 bringup、标定、诊断、日志、安全策略和自动化测试"),
], [1.55, 4.72])
add_heading(doc, "项目仓库", 2)
p = add_para(doc, "", after=4)
add_hyperlink(p, "https://github.com/2799063570/wheeltec_robot", "https://github.com/2799063570/wheeltec_robot")
add_para(doc, "建议在正式投递前补充 30–90 秒演示视频，并在 GitHub README 中加入 Gazebo、RViz、视觉识别和抓取分拣截图。", size=9.6, color=MID_GRAY)
add_callout(doc, "投递前最后检查", "请填写封面的姓名/学校/团队信息，并根据实际参与情况修改“个人贡献”部分；如有实物演示、比赛奖项或量化测试结果，可在本页追加。", YELLOW, ORANGE)

# Document properties.
doc.core_properties.title = "基于 ROS 的 AUBO 复合移动机器人项目介绍"
doc.core_properties.subject = "自主导航、视觉识别、机械臂抓取与颜色分拣"
doc.core_properties.author = "项目团队（请补充）"
doc.core_properties.keywords = "ROS, AUBO, 移动机器人, 自主导航, MoveIt, OpenCV, 视觉分拣"

doc.save(OUT)
print(OUT)
